import io
import re
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple

from docx import Document


@dataclass
class TranscriptUtterance:
    number: int
    raw: str
    normalized: str


@dataclass
class TranscriptData:
    meta: Dict[str, str]
    utterances: List[TranscriptUtterance]


_RE_DATE = re.compile(r"(\d{1,2})\s*/\s*(\d{1,2})\s*/\s*(\d{2,4})")


def _normalize_text_for_model(text: str) -> str:
    """A pragmatic normalization: keep content, remove formatting artifacts."""
    s = (text or "").strip()
    if not s:
        return ""

    # Quotes / underscores
    s = s.replace("“", "").replace("”", "").replace('"', "")
    s = s.replace("_", "")

    # Join hard hyphen breaks like "come- dor" -> "comedor"
    s = re.sub(r"(\w)\s*-\s*(\w)", r"\1\2", s)

    # Treat slash-delimited items as space-delimited for readability
    s = s.replace("/", " ")

    # Remove stray tokens that are just punctuation
    s = re.sub(r"\b[¿?¡!]+\b", " ", s)

    # Keep Spanish question marks if they wrap a word (e.g., ¿Mentira?)
    # but remove other punctuation.
    s = re.sub(r"[^\w\sáéíóúüñÁÉÍÓÚÜÑ¿?¡!]", " ", s)

    # Collapse whitespace
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _extract_meta_from_paragraphs(paragraphs: List[str]) -> Dict[str, str]:
    meta: Dict[str, str] = {}
    for p in paragraphs:
        p_clean = p.strip()
        if not p_clean:
            continue

        if "Nombre del aprendiz" in p_clean:
            meta["learner_name"] = p_clean.split(":", 1)[-1].strip()
        elif p_clean.lower().startswith("fecha"):
            meta["date_raw"] = p_clean.split(":", 1)[-1].strip()
            m = _RE_DATE.search(meta["date_raw"])
            if m:
                dd, mm, yy = m.groups()
                if len(yy) == 2:
                    yy = "20" + yy
                meta["date_iso"] = f"{yy.zfill(4)}-{mm.zfill(2)}-{dd.zfill(2)}"
        elif "Sesion" in p_clean or "Sesión" in p_clean:
            meta["session"] = p_clean.split(":", 1)[-1].strip()
        elif "Muestra" in p_clean:
            meta["sample"] = p_clean.split(":", 1)[-1].strip()

    return meta


def _parse_tables_with_continuations(doc: Document) -> List[Tuple[int, str]]:
    """Extract (utterance_number, utterance_text) from transcript-like tables.

    Handles multi-table layouts and continuation rows where the number cell is blank.
    """
    utterances: Dict[int, str] = {}
    last_num: Optional[int] = None
    last_utt_idx: Optional[int] = None

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]

            # Find the first purely-numeric cell in the row
            num = None
            num_idx = None
            for idx, val in enumerate(cells):
                if re.fullmatch(r"\d+", val or ""):
                    num = int(val)
                    num_idx = idx
                    break

            if num is not None and num_idx is not None:
                # Next non-empty cell is typically the utterance
                utt = ""
                utt_idx = None
                for j in range(num_idx + 1, len(cells)):
                    if cells[j]:
                        utt = cells[j]
                        utt_idx = j
                        break

                if utt and utt_idx is not None:
                    utterances[num] = utt
                    last_num = num
                    last_utt_idx = utt_idx
                continue

            # Continuation row: append text in the previous utterance column
            if last_num is not None and last_utt_idx is not None and last_utt_idx < len(cells):
                cont = (cells[last_utt_idx] or "").strip()
                if cont:
                    utterances[last_num] = (utterances.get(last_num, "") + " " + cont).strip()

    return [(n, utterances[n]) for n in sorted(utterances.keys())]


def parse_transcript_docx(docx_bytes: bytes) -> TranscriptData:
    doc = Document(io.BytesIO(docx_bytes))

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    meta = _extract_meta_from_paragraphs(paragraphs)

    pairs = _parse_tables_with_continuations(doc)

    utterances: List[TranscriptUtterance] = []
    for n, raw in pairs:
        norm = _normalize_text_for_model(raw)
        if norm:
            utterances.append(TranscriptUtterance(number=n, raw=raw, normalized=norm))

    return TranscriptData(meta=meta, utterances=utterances)


def parse_transcript_txt(text: str) -> TranscriptData:
    """Accepts either:
    - raw transcript with numbered lines (e.g., '1. ...')
    - or unnumbered text; we will treat each non-empty line as an utterance.
    """
    lines = [ln.strip() for ln in (text or "").splitlines()]
    lines = [ln for ln in lines if ln]

    utterances: List[TranscriptUtterance] = []

    for ln in lines:
        m = re.match(r"^(\d+)\s*[\.)-]\s*(.+)$", ln)
        if m:
            n = int(m.group(1))
            raw = m.group(2).strip()
        else:
            n = (utterances[-1].number + 1) if utterances else 1
            raw = ln

        norm = _normalize_text_for_model(raw)
        if norm:
            utterances.append(TranscriptUtterance(number=n, raw=raw, normalized=norm))

    return TranscriptData(meta={}, utterances=utterances)


def build_numbered_transcript_block(utterances: List[TranscriptUtterance]) -> str:
    return "\n".join([f"{u.number}. {u.normalized}" for u in utterances])
