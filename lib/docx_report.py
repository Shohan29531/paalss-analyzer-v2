import io
import re
from typing import Optional

from docx import Document
from docx.shared import Pt


def report_text_to_docx_bytes(report_text: str, title: Optional[str] = None) -> bytes:
    """Convert a plain-text PAALSS report into a simple .docx."""
    doc = Document()

    # A slightly nicer default font size
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    text = (report_text or "").strip("\n")
    if not text:
        doc.add_paragraph("(empty report)")
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()

    lines = text.splitlines()

    # Optional explicit title override
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(16)
        doc.add_paragraph("")

    for ln in lines:
        if not ln.strip():
            doc.add_paragraph("")
            continue

        # Heuristics for headings
        is_main_title = ln.strip().lower() == "paalss comprehensive language sample report"
        is_section = bool(re.match(r"^\d+\.\s+", ln.strip()))
        is_subhead = ln.strip().endswith(":") and len(ln.strip()) <= 80

        if is_main_title:
            p = doc.add_paragraph()
            r = p.add_run(ln.strip())
            r.bold = True
            r.font.size = Pt(16)
        elif is_section:
            p = doc.add_paragraph()
            r = p.add_run(ln.strip())
            r.bold = True
            r.font.size = Pt(13)
        elif is_subhead:
            p = doc.add_paragraph()
            r = p.add_run(ln.strip())
            r.bold = True
        else:
            doc.add_paragraph(ln)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
